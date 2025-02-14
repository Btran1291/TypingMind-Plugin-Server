from flask import Flask, request, jsonify, send_file, url_for, Blueprint
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_UNDERLINE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_DIRECTION
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
import io
import uuid
import os
import requests
from datetime import datetime

generate_docx_bp = Blueprint('generate_docx', __name__)
CORS(generate_docx_bp, resources={r"/*": {"origins": "*"}})
generated_files = {}

def process_core_properties(core_properties_data, core_properties):
    if core_properties_data:
        for prop, value in core_properties_data.items():
            if hasattr(core_properties, prop):
                if prop in ['created', 'modified', 'last_printed']:
                    try:
                        setattr(core_properties, prop, datetime.fromisoformat(value))
                    except (ValueError, TypeError):
                        print(f"Invalid date format for {prop}: {value}. Skipping.")
                    except Exception as e:
                        print(f"Error setting core property {prop}: {e}")
                else:
                    setattr(core_properties, prop, value)

def process_headers_footers(section, headers_data, footers_data):
    def apply_paragraph_formatting(paragraph, format_data):
        if format_data:
            p_format = paragraph.paragraph_format
            for key, value in format_data.items():
                if key == 'alignment':
                    p_format.alignment = getattr(WD_ALIGN_PARAGRAPH, value, WD_ALIGN_PARAGRAPH.LEFT)
                elif key == 'line_spacing':
                    p_format.line_spacing = value
                elif key == 'line_spacing_rule':
                    p_format.line_spacing_rule = getattr(WD_LINE_SPACING, value, WD_LINE_SPACING.SINGLE)
                elif key == 'first_line_indent':
                    p_format.first_line_indent = Inches(value)
                elif key == 'left_indent':
                    p_format.left_indent = Inches(value)
                elif key == 'right_indent':
                    p_format.right_indent = Inches(value)
                elif key == 'space_before':
                    p_format.space_before = Pt(value)
                elif key == 'space_after':
                    p_format.space_after = Pt(value)
                elif key == 'keep_together':
                    p_format.keep_together = value
                elif key == 'keep_with_next':
                    p_format.keep_with_next = value
                elif key == 'page_break_before':
                    p_format.page_break_before = value
                elif key == 'widow_control':
                    p_format.widow_control = value
                elif key == 'tab_stops':
                    for tab_stop in value:
                        p_format.tab_stops.add_tab_stop(
                            Inches(tab_stop.get('position', 0)),
                            alignment=getattr(WD_TAB_ALIGNMENT, tab_stop.get('alignment', 'LEFT'), WD_TAB_ALIGNMENT.LEFT),
                            leader=getattr(WD_TAB_LEADER, tab_stop.get('leader', 'SPACES'), WD_TAB_LEADER.SPACES)
                        )
    def apply_font_formatting(paragraph, font_data):
        if font_data:
            font = paragraph.runs[0].font if paragraph.runs else paragraph.add_run().font
            for font_key, font_value in font_data.items():
                if font_key == 'color':
                    if isinstance(font_value, str) and font_value.startswith('#'):
                        try:
                            rgb = RGBColor.from_string(font_value[1:])
                            font.color.rgb = rgb
                        except ValueError:
                            print(f"Invalid color format: {font_value}. Skipping.")
                    else:
                        print(f"Invalid color format: {font_value}. Skipping.")
                elif font_key == 'size':
                    font.size = Pt(font_value)
                elif hasattr(font, font_key):
                    setattr(font, font_key, font_value)

    if headers_data:
        for header_type, header_content in headers_data.items():
            if header_type == "FIRST":
                header = section.first_page_header
            elif header_type == "EVEN":
                header = section.even_page_header
            else:
                header = section.header

            if header:
                header.is_linked_to_previous = False
                if header.paragraphs:
                    paragraph = header.paragraphs[0]
                    paragraph.text = header_content
                else:
                    paragraph = header.add_paragraph(header_content)

                apply_paragraph_formatting(paragraph, headers_data.get('paragraph_format'))
                if 'font' in headers_data:
                    apply_font_formatting(paragraph, headers_data['font'])

    if footers_data:
        for footer_type, footer_content in footers_data.items():
            if footer_type == "FIRST":
                footer = section.first_page_footer
            elif footer_type == "EVEN":
                footer = section.even_page_footer
            else:
                footer = section.footer

            if footer:
                footer.is_linked_to_previous = False
                if footer.paragraphs:
                    paragraph = footer.paragraphs[0]
                    paragraph.text = footer_content
                else:
                    paragraph = footer.add_paragraph(footer_content)
                if 'paragraph_format' in footers_data:
                    apply_paragraph_formatting(paragraph, footers_data['paragraph_format'])
                if 'font' in footers_data:
                    apply_font_formatting(paragraph, footers_data['font'])

def process_sections(document, sections_data, default_page_width, default_page_height,
                     default_left_margin, default_right_margin, default_top_margin,
                     default_bottom_margin, default_gutter, default_header_distance,
                     default_footer_distance, default_orientation):
    for section_data in sections_data:
        section = document.add_section(start_type=getattr(WD_SECTION_START, section_data.get('start_type', 'NEW_PAGE'), WD_SECTION_START.NEW_PAGE))
        section.orientation = getattr(WD_ORIENTATION, section_data.get('orientation', default_orientation), WD_ORIENTATION.PORTRAIT)
        section.page_width = Inches(section_data.get('page_width', default_page_width))
        section.page_height = Inches(section_data.get('page_height', default_page_height))
        section.left_margin = Inches(section_data.get('left_margin', default_left_margin))
        section.right_margin = Inches(section_data.get('right_margin', default_right_margin))
        section.top_margin = Inches(section_data.get('top_margin', default_top_margin))
        section.bottom_margin = Inches(section_data.get('bottom_margin', default_bottom_margin))
        section.gutter = Inches(section_data.get('gutter', default_gutter))
        section.header_distance = Inches(section_data.get('header_distance', default_header_distance))
        section.footer_distance = Inches(section_data.get('footer_distance', default_footer_distance))
        process_headers_footers(section, section_data.get('headers'), section_data.get('footers'))

def process_paragraph(document, item):
    style_name = item.get('style')
    style = None
    if style_name:
        try:
            style = document.styles[style_name]
        except KeyError:
            print(f"Style '{style_name}' not found. Using default style.")
    paragraph = document.add_paragraph(item.get('text', ''), style=style)
    if 'paragraph_format' in item:
        p_format = paragraph.paragraph_format
        for key, value in item['paragraph_format'].items():
            if key == 'alignment':
                p_format.alignment = getattr(WD_ALIGN_PARAGRAPH, value, WD_ALIGN_PARAGRAPH.LEFT)
            elif key == 'line_spacing':
                p_format.line_spacing = value
            elif key == 'line_spacing_rule':
                p_format.line_spacing_rule = getattr(WD_LINE_SPACING, value, WD_LINE_SPACING.SINGLE)
            elif key == 'first_line_indent':
                p_format.first_line_indent = Inches(value)
            elif key == 'left_indent':
                p_format.left_indent = Inches(value)
            elif key == 'right_indent':
                p_format.right_indent = Inches(value)
            elif key == 'space_before':
                p_format.space_before = Pt(value)
            elif key == 'space_after':
                p_format.space_after = Pt(value)
            elif key == 'keep_together':
                p_format.keep_together = value
            elif key == 'keep_with_next':
                p_format.keep_with_next = value
            elif key == 'page_break_before':
                p_format.page_break_before = value
            elif key == 'widow_control':
                p_format.widow_control = value
            elif key == 'tab_stops':
                for tab_stop in value:
                    p_format.tab_stops.add_tab_stop(
                        Inches(tab_stop.get('position', 0)),
                        alignment=getattr(WD_TAB_ALIGNMENT, tab_stop.get('alignment', 'LEFT'), WD_TAB_ALIGNMENT.LEFT),
                        leader=getattr(WD_TAB_LEADER, tab_stop.get('leader', 'SPACES'), WD_TAB_LEADER.SPACES)
                    )
    if 'runs' in item:
        for run_data in item['runs']:
            run = paragraph.add_run(run_data.get('text', ''))
            if 'font' in run_data:
                font = run.font
                for font_key, font_value in run_data['font'].items():
                    if font_key == 'color':
                        if isinstance(font_value, str) and font_value.startswith('#'):
                            try:
                                rgb = RGBColor.from_string(font_value[1:])
                                font.color.rgb = rgb
                            except ValueError:
                                print(f"Invalid color format: {font_value}. Skipping.")
                        else:
                            print(f"Invalid color format: {font_value}. Skipping.")
                    elif font_key == 'size':
                        font.size = Pt(font_value)
                    elif hasattr(font, font_key):
                        setattr(font, font_key, font_value)
            if 'bold' in run_data:
                run.bold = run_data['bold']
            if 'italic' in run_data:
                run.italic = run_data['italic']
            if 'underline' in run_data:
                run.underline = getattr(WD_UNDERLINE, run_data['underline'], WD_UNDERLINE.NONE)

def process_table(document, item):
    rows = item.get('rows', 1)
    cols = item.get('cols', 1)
    style_name = item.get('style')
    style = None
    if style_name:
        try:
            style = document.styles[style_name]
        except KeyError:
            print(f"Style '{style_name}' not found. Using default style.")
    table = document.add_table(rows=rows, cols=cols, style=style)
    if 'alignment' in item:
        table.alignment = getattr(WD_TABLE_ALIGNMENT, item['alignment'], WD_TABLE_ALIGNMENT.LEFT)
    if 'table_direction' in item:
        table.table_direction = getattr(WD_TABLE_DIRECTION, item['table_direction'], WD_TABLE_DIRECTION.LTR)
    if 'autofit' in item:
        table.autofit = item['autofit']
    if 'data' in item:
        for row_idx, row_data in enumerate(item['data']):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = str(cell_data)
                if 'vertical_alignment' in item:
                    cell.vertical_alignment = getattr(WD_CELL_VERTICAL_ALIGNMENT, item['vertical_alignment'], WD_CELL_VERTICAL_ALIGNMENT.TOP)

def process_image(document, item):
    try:
        headers = {"User-Agent": f"Docx_Generator_bot/1.0 requests/{requests.__version__}"}
        response = requests.get(item['url'], headers=headers, stream=True)
        response.raise_for_status()
        image_stream = io.BytesIO(response.content)
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        width = Inches(item.get('width', 3))
        height = Inches(item.get('height', 2))
        run.add_picture(image_stream, width=width, height=height)
    except Exception as e:
        print(f"Error adding image: {e}")

@generate_docx_bp.route('/generate_docx', methods=['POST', 'OPTIONS'])
def generate_docx():
    if request.method == 'OPTIONS':
        return jsonify({'status': 'OK'}), 200

    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'Invalid input. Must provide document parameters.'}), 400

        default_page_width = data.get('defaultPageWidth', 8.5)
        default_page_height = data.get('defaultPageHeight', 11)
        default_left_margin = data.get('defaultLeftMargin', 1)
        default_right_margin = data.get('defaultRightMargin', 1)
        default_top_margin = data.get('defaultTopMargin', 1)
        default_bottom_margin = data.get('defaultBottomMargin', 1)
        default_gutter = data.get('defaultGutter', 0)
        default_header_distance = data.get('defaultHeaderDistance', 0.5)
        default_footer_distance = data.get('defaultFooterDistance', 0.5)
        default_orientation = data.get('defaultOrientation', "PORTRAIT")
        enable_core_properties = data.get('enableCoreProperties', "false") == "true"
        core_properties_title = data.get('corePropertiesTitle', "")
        core_properties_author = data.get('corePropertiesAuthor', "")
        core_properties_created = data.get('corePropertiesCreated', "")
        odd_and_even_pages_header_footer = data.get('oddAndEvenPagesHeaderFooter', "false") == "true"

        document = Document()
        settings = document.settings
        settings.odd_and_even_pages_header_footer = odd_and_even_pages_header_footer

        if enable_core_properties:
            core_properties = document.core_properties
            core_properties_data = {}
            if core_properties_title:
                core_properties_data['title'] = core_properties_title
            if core_properties_author:
                core_properties_data['author'] = core_properties_author
            if core_properties_created:
                core_properties_data['created'] = core_properties_created
            process_core_properties(core_properties_data, core_properties)

        process_sections(document, data.get('sections', []), default_page_width, default_page_height,
                         default_left_margin, default_right_margin, default_top_margin,
                         default_bottom_margin, default_gutter, default_header_distance,
                         default_footer_distance, default_orientation)

        content = data.get('content', [])
        for item in content:
            if item['type'] == 'heading':
                document.add_heading(item.get('text', ''), level=item.get('level', 1))
            elif item['type'] == 'paragraph':
                process_paragraph(document, item)
            elif item['type'] == 'table':
                process_table(document, item)
            elif item['type'] == 'image':
                process_image(document, item)
            elif item['type'] == 'list':
                list_style = item.get('numbering_style') or item.get('style', 'List Bullet')
                for list_item in item.get('items', []):
                    document.add_paragraph(list_item, style=list_style)
            elif item['type'] == 'page_break':
                document.add_page_break()

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        filename = f"document_{uuid.uuid4()}.docx"
        temp_file_path = os.path.join("/tmp", filename)
        with open(temp_file_path, 'wb') as f:
            f.write(buffer.read())

        file_id = str(uuid.uuid4())
        generated_files[file_id] = temp_file_path

        download_link = url_for('generate_docx.download_file', file_id=file_id, _external=True)

        return jsonify({'download_link': download_link})

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@generate_docx_bp.route('/download/<file_id>')
def download_file(file_id):
    if file_id in generated_files:
        temp_file_path = generated_files[file_id]
        return send_file(
            temp_file_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            download_name='generated_document.docx',
            as_attachment=True
        )
    else:
        return "File not found", 404
